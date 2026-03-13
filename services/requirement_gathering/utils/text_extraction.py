"""
Text extraction utility for various document formats.
Supports: PDF, DOCX, TXT files
"""
from pathlib import Path
from typing import Optional
import re


def extract_text_from_txt(file_path: Path) -> str:
    """Extract text from TXT file"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    except UnicodeDecodeError:
        # Try with different encoding
        with open(file_path, 'r', encoding='latin-1') as f:
            return f.read()


def extract_text_from_pdf(file_path: Path) -> str:
    """Extract text from PDF file using PyPDF2"""
    try:
        import PyPDF2
        
        text = []
        with open(file_path, 'rb') as f:
            pdf_reader = PyPDF2.PdfReader(f)
            for page_num, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text()
                if page_text.strip():
                    text.append(f"--- Page {page_num + 1} ---\n{page_text}")
        
        return "\n\n".join(text)
    except ImportError:
        raise Exception("PyPDF2 not installed. Install with: pip install PyPDF2")
    except Exception as e:
        raise Exception(f"Failed to extract text from PDF: {str(e)}")


def extract_text_from_docx(file_path: Path) -> str:
    """Extract text from DOCX file using python-docx"""
    try:
        from docx import Document
        
        doc = Document(file_path)
        text = []
        
        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text.append(para.text)
        
        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text.append(row_text)
        
        return "\n\n".join(text)
    except ImportError:
        raise Exception("python-docx not installed. Install with: pip install python-docx")
    except Exception as e:
        raise Exception(f"Failed to extract text from DOCX: {str(e)}")


def extract_text(file_path: Path) -> tuple[str, str]:
    """
    Extract text from supported document formats.
    
    Args:
        file_path: Path to the file
        
    Returns:
        Tuple of (extracted_text, file_type)
        
    Raises:
        ValueError: If file format is not supported
        Exception: If extraction fails
    """
    if not file_path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")
    
    suffix = file_path.suffix.lower()
    
    extractors = {
        '.txt': ('txt', extract_text_from_txt),
        '.pdf': ('pdf', extract_text_from_pdf),
        '.docx': ('docx', extract_text_from_docx),
        '.doc': ('doc', extract_text_from_docx),  # Try docx extractor for .doc
    }
    
    if suffix not in extractors:
        raise ValueError(f"Unsupported file format: {suffix}. Supported: {', '.join(extractors.keys())}")
    
    file_type, extractor = extractors[suffix]
    
    try:
        text = extractor(file_path)
        
        # Clean up text
        text = clean_text(text)
        
        if not text or len(text.strip()) < 10:
            raise Exception("Extracted text is too short or empty")
        
        return text, file_type
        
    except Exception as e:
        raise Exception(f"Text extraction failed for {file_path.name}: {str(e)}")


def clean_text(text: str) -> str:
    """Clean and normalize extracted text"""
    # Remove excessive whitespace
    text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
    text = re.sub(r' +', ' ', text)
    
    # Remove control characters except newlines and tabs
    text = ''.join(char for char in text if char.isprintable() or char in '\n\t')
    
    return text.strip()


def chunk_text(text: str, chunk_size: int = 1000, overlap: int = 200) -> list[dict]:
    """
    Split text into overlapping chunks with metadata.
    
    Args:
        text: Text to chunk
        chunk_size: Target size of each chunk (characters)
        overlap: Overlap between consecutive chunks (characters)
        
    Returns:
        List of chunk dictionaries with text and metadata
    """
    # Split by paragraphs first
    paragraphs = text.split('\n\n')
    
    chunks = []
    current_chunk = []
    current_size = 0
    chunk_index = 0
    
    for para in paragraphs:
        para = para.strip()
        if not para:
            continue
        
        para_size = len(para)
        
        # If single paragraph exceeds chunk_size, split it
        if para_size > chunk_size:
            # Save current chunk if exists
            if current_chunk:
                chunk_text = '\n\n'.join(current_chunk)
                chunks.append({
                    'text': chunk_text,
                    'chunk_index': chunk_index,
                    'char_count': len(chunk_text)
                })
                chunk_index += 1
                current_chunk = []
                current_size = 0
            
            # Split large paragraph by sentences
            sentences = re.split(r'(?<=[.!?])\s+', para)
            for sentence in sentences:
                if len(sentence) > chunk_size:
                    # Split very long sentences by words
                    words = sentence.split()
                    temp_chunk = []
                    temp_size = 0
                    
                    for word in words:
                        word_size = len(word) + 1
                        if temp_size + word_size > chunk_size and temp_chunk:
                            chunk_text = ' '.join(temp_chunk)
                            chunks.append({
                                'text': chunk_text,
                                'chunk_index': chunk_index,
                                'char_count': len(chunk_text)
                            })
                            chunk_index += 1
                            temp_chunk = []
                            temp_size = 0
                        
                        temp_chunk.append(word)
                        temp_size += word_size
                    
                    if temp_chunk:
                        current_chunk.append(' '.join(temp_chunk))
                        current_size += temp_size
                else:
                    current_chunk.append(sentence)
                    current_size += len(sentence) + 2
        
        # Add paragraph to current chunk if it fits
        elif current_size + para_size + 2 <= chunk_size:
            current_chunk.append(para)
            current_size += para_size + 2
        else:
            # Save current chunk and start new one
            if current_chunk:
                chunk_text = '\n\n'.join(current_chunk)
                chunks.append({
                    'text': chunk_text,
                    'chunk_index': chunk_index,
                    'char_count': len(chunk_text)
                })
                chunk_index += 1
            
            current_chunk = [para]
            current_size = para_size
    
    # Add final chunk
    if current_chunk:
        chunk_text = '\n\n'.join(current_chunk)
        chunks.append({
            'text': chunk_text,
            'chunk_index': chunk_index,
            'char_count': len(chunk_text)
        })
    
    return chunks
